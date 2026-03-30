#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将机器人数据收集进展报告转换为 Word 文档
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

# 读取 markdown 文件
with open('/home/admin/.openclaw/workspace/机器人数据收集进展报告.md', 'r', encoding='utf-8') as f:
    content = f.read()

# 创建 Word 文档
doc = Document()

# 设置样式
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(11)

# 添加标题
title = doc.add_heading('机器人数据库收集进展报告', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 添加版本信息
doc.add_paragraph('生成时间：2026-03-20 11:45\n数据版本：v7.0')

# 解析 markdown 内容
lines = content.split('\n')
current_section = None

for line in lines:
    line = line.strip()
    
    # 跳过空行和分隔线
    if not line or line.startswith('---'):
        continue
    
    # 处理标题
    if line.startswith('## '):
        section_title = line.replace('## ', '').replace('**', '')
        doc.add_heading(section_title, level=1)
    elif line.startswith('### '):
        subsection_title = line.replace('### ', '').replace('**', '')
        doc.add_heading(subsection_title, level=2)
    elif line.startswith('#### '):
        subsubsection_title = line.replace('#### ', '').replace('**', '')
        doc.add_heading(subsubsection_title, level=3)
    
    # 处理列表项
    elif line.startswith('- **') or line.startswith('- ⏳') or line.startswith('- ✅'):
        # 清理格式
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
        text = text.replace('- ', '').strip()
        doc.add_paragraph(text, style='List Bullet')
    
    # 处理表格行
    elif line.startswith('|') and line.endswith('|'):
        # 简单处理表格，转换为文本
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        if cells and cells[0] not in ['------', '公司', '领域', '文件名']:
            table_text = '  '.join(cells)
            doc.add_paragraph(table_text)
    
    # 处理普通文本
    elif not line.startswith('#') and not line.startswith('|'):
        # 清理 markdown 格式
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
        text = re.sub(r'`([^`]+)`', r'\1', text)
        if text and len(text) > 2:
            doc.add_paragraph(text)

# 添加页脚信息
doc.add_paragraph('\n' + '='*50)
doc.add_paragraph('报告更新时间：2026-03-20 11:45')
doc.add_paragraph('数据范围：174 款产品 / 72 起事故 / 47 款销售数据')

# 保存 Word 文档
output_path = '/home/admin/.openclaw/workspace/机器人数据收集进展报告.docx'
doc.save(output_path)

print(f"Word 文档已生成：{output_path}")
print(f"文件大小：{len(content)} 字符")
