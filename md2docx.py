#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Markdown to Word DOCX converter"""

import sys
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def read_markdown(filepath):
    with open(filepath, 'r', encoding='utf-8') as f:
        return f.read()

def parse_markdown_table(table_text):
    """解析 Markdown 表格"""
    lines = table_text.strip().split('\n')
    if len(lines) < 3:
        return []
    
    rows = []
    for line in lines:
        if '|---' in line:
            continue
        cells = [cell.strip() for cell in line.split('|')]
        if cells and cells[0] == '':
            cells = cells[1:]
        if cells and cells[-1] == '':
            cells = cells[:-1]
        if cells:
            rows.append(cells)
    return rows

def convert_md_to_docx(md_content, output_path):
    doc = Document()
    
    # 设置中文字体支持
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(10.5)
    
    lines = md_content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # 处理标题
        if line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
        # 处理表格
        elif line.startswith('|') and '|---' not in line:
            table_lines = [line]
            i += 1
            while i < len(lines) and lines[i].startswith('|'):
                if '|---' in lines[i]:
                    i += 1
                    continue
                table_lines.append(lines[i])
                i += 1
            i -= 1
            
            table_data = parse_markdown_table('\n'.join(table_lines))
            if table_data:
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Table Grid'
                for row_idx, row_data in enumerate(table_data):
                    for col_idx, cell_data in enumerate(row_data):
                        cell = table.cell(row_idx, col_idx)
                        cell.text = cell_data
                        if row_idx == 0:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
        # 处理粗体
        elif line.startswith('**') and line.endswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(line[2:-2])
            run.bold = True
        # 处理列表
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        # 处理空行
        elif line.strip() == '':
            doc.add_paragraph('')
        # 普通段落
        else:
            # 处理行内粗体
            if '**' in line:
                p = doc.add_paragraph()
                parts = re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        p.add_run(part)
            else:
                doc.add_paragraph(line)
        
        i += 1
    
    doc.save(output_path)
    print(f"已保存：{output_path}")

if __name__ == '__main__':
    if len(sys.argv) != 3:
        print("用法：python md2docx.py <input.md> <output.docx>")
        sys.exit(1)
    
    md_content = read_markdown(sys.argv[1])
    convert_md_to_docx(md_content, sys.argv[2])
